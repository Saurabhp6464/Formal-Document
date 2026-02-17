from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from django.conf import settings
from django.utils import timezone
from django.template.loader import render_to_string
from django.contrib.auth import authenticate, login as auth_login, logout as auth_logout
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.contrib import messages

from io import BytesIO
import json, os
from datetime import datetime

from .models import OfficeOrderCounter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from weasyprint import HTML, CSS

import google.generativeai as genai

from .constants import DESIGNATION_MAP

import uuid


# DOCUMENT SELECTOR - REMOVED


# ---------------- GEMINI ----------------
genai.configure(api_key=settings.GEMINI_API_KEY)
gemini_model = genai.GenerativeModel("models/gemini-2.5-flash")

# ---------------- LOAD JSON ----------------
BASE_DIR = settings.BASE_DIR
with open(os.path.join(BASE_DIR, "office_order.json"), encoding="utf-8") as f:
    OFFICE_ORDER = json.load(f)

with open(os.path.join(BASE_DIR, "circular.json"), encoding="utf-8") as f:
    CIRCULAR = json.load(f)

with open(os.path.join(BASE_DIR, "policy.json"), encoding="utf-8") as f:
    POLICY = json.load(f)

# ---------------- HELPERS ----------------
def format_date_ddmmyyyy(date_str):
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d-%m-%Y")
    except Exception:
        return date_str

# ---------------- HOME ----------------
@login_required(login_url='login')
def home(request):
    return render(request, "generator/home.html", {
        "designations": DESIGNATION_MAP.keys(),
        "people": CIRCULAR["people"]
    })


# ===================== OFFICE ORDER (UNCHANGED) ======================


def generate_body(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    try:
        prompt = request.POST.get("body_prompt", "").strip()
        lang = request.POST.get("language", "en")

        if not prompt:
            return HttpResponse("Please provide a prompt for AI generation.", status=200)

        if lang == "hi":
            system_prompt = """
आप BISAG-N के लिए एक आधिकारिक कार्यालय आदेश की मुख्य सामग्री लिख रहे हैं।

नियम:
- कम से कम 2–3 वाक्यों का एक औपचारिक अनुच्छेद लिखें।
- सरकारी भाषा का प्रयोग करें।
- कोई शीर्षक, संदर्भ, दिनांक, प्रेषक या प्राप्तकर्ता न लिखें।
- बुलेट या क्रमांक का प्रयोग न करें।
- केवल सादा पाठ में उत्तर दें।
"""
        else:
            system_prompt = """
You are drafting the BODY of an official government Office Order for BISAG-N.

Rules:
- Write one formal paragraph (minimum 2–3 sentences).
- Use official government tone.
- Do not include title, reference, date, From or To.
- No bullet points or numbering.
- Plain text only.
"""

        res = gemini_model.generate_content(system_prompt + "\n\nTopic:\n" + prompt)
        generated_text = res.text.strip() if res and res.text else "Unable to generate content. Please try again."
        return HttpResponse(generated_text, status=200)
        
    except Exception as e:
        error_msg = f"AI generation failed: {str(e)}"
        print(f"[ERROR] Gemini API: {error_msg}")
        return HttpResponse("Failed to generate content. Please check your internet connection and try again.", status=200)


def result_office_order(request):
    if request.method != "POST":
        return redirect("home")

    lang = request.POST.get("language", "en")
    raw_date = request.POST.get("date")
    date = format_date_ddmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%m-%Y")

    # Auto-generate reference number with sequential counter
    current_year = timezone.now().year
    counter = OfficeOrderCounter.get_next_number(current_year)
    
    if lang == "hi":
        reference = f"बायसेग-एन/कार्यालय आदेश/{current_year}/{counter:03d}"
    else:
        reference = f"BISAG-N/Office Order/{current_year}/{counter:03d}"

    data = {
        "language": lang,
        "header": OFFICE_ORDER["header"][lang],
        "title": OFFICE_ORDER["title_hi"] if lang == "hi" else OFFICE_ORDER["title_en"],
        "reference": reference,
        "date": date,
        "body": request.POST.get("body", "").strip(),
        "from": DESIGNATION_MAP[request.POST.get("from_position")][lang],
        "to": [DESIGNATION_MAP[x][lang] for x in request.POST.getlist("to_recipients[]")],
    }

    request.session["doc_data"] = data
    return render(request, "generator/result_office_order.html", data)

# PDF + DOCX for office order → UNCHANGED
# (your existing download_pdf & download_docx remain exactly same)

# ========================= CIRCULAR (UPDATED) ========================


def circular_form(request):
    return render(request, "generator/circular_form.html", {
        "people": CIRCULAR["people"]
    })

# -------- GEMINI CIRCULAR BODY --------
def generate_circular_body(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    try:
        prompt = request.POST.get("body_prompt", "").strip()
        lang = request.POST.get("language", "en")
        
        if not prompt:
            return HttpResponse("Please provide a prompt for AI generation.", status=200)

        if lang == "hi":
            system_prompt = """
आप BISAG-N के लिए एक सरकारी परिपत्र (Circular) का केवल मुख्य भाग (BODY) लिख रहे हैं।

महत्वपूर्ण नियम:
- केवल परिपत्र का मुख्य विषय-वस्तु लिखें।
- कोई विषय (Subject) न लिखें।
- कोई शीर्षक न लिखें।
- कोई संदर्भ संख्या न लिखें।
- कोई हस्ताक्षर न लिखें।
- कोई दिनांक न लिखें।
- कोई "प्रेषक" या "प्राप्तकर्ता" न लिखें।
- 1–2 औपचारिक अनुच्छेद लिखें।
- सरकारी भाषा का प्रयोग करें।
- केवल सादा पाठ में उत्तर दें।
"""
        else:
            system_prompt = """
You are drafting ONLY the BODY content of an official Government Circular for BISAG-N.

IMPORTANT Rules:
- Write ONLY the main body content of the circular.
- Do NOT include any subject line.
- Do NOT include any title or heading.
- Do NOT include reference number.
- Do NOT include signature.
- Do NOT include date.
- Do NOT include From or To sections.
- Write 1–2 formal paragraphs only.
- Official government tone.
- Plain text only.
"""

        res = gemini_model.generate_content(system_prompt + "\n\nTopic:\n" + prompt)
        generated_text = res.text.strip() if res and res.text else "Unable to generate content. Please try again."
        return HttpResponse(generated_text, status=200)
        
    except Exception as e:
        error_msg = f"AI generation failed: {str(e)}"
        print(f"[ERROR] Circular Gemini API: {error_msg}")
        return HttpResponse("Failed to generate content. Please check your internet connection and try again.", status=200)

# -------- CIRCULAR PREVIEW --------
def result_circular(request):
    if request.method != "POST":
        return redirect("circular_form")

    lang = request.POST.get("language")
    raw_date = request.POST.get("date")
    date = format_date_ddmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%m-%Y")
    subject = request.POST.get("subject")
    body = request.POST.get("body")
    
    # Get from position from dropdown (like office order)
    from_position = request.POST.get("from_position")
    from_designation = DESIGNATION_MAP[from_position][lang] if from_position else ""
    
    to_ids = request.POST.getlist("to[]")

    people = CIRCULAR["people"]
    to_people = [p for p in people if str(p["id"]) in to_ids]

    # Get header based on language
    if lang == "hi":
        header = {
            "org_name": CIRCULAR["header"]["hindi"]["org_name"],
            "ministry": CIRCULAR["header"]["hindi"]["ministry"],
            "government": CIRCULAR["header"]["hindi"]["government"]
        }
    else:
        header = {
            "org_name": CIRCULAR["header"]["english"]["org_name"],
            "ministry": CIRCULAR["header"]["english"]["ministry"],
            "government": CIRCULAR["header"]["english"]["government"]
        }

    data = {
        "language": lang,
        "header": header,
        "date": date,
        "subject": subject,
        "body": body,
        "from": from_designation,
        "to_people": to_people,
    }

    request.session["circular_data"] = data
    return render(request, "generator/result_circular.html", data)

# -------- CIRCULAR PDF --------
def download_circular_pdf(request):
    data = request.session.get("circular_data")
    if not data:
        return HttpResponse("No circular generated", status=400)

    html = render_to_string("generator/pdf_circular.html", data)

    # Optimize PDF generation with font subsetting and compression
    pdf = HTML(
        string=html,
        base_url=settings.BASE_DIR
    ).write_pdf(
        optimize_images=True,
        jpeg_quality=85,
        presentational_hints=True
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Circular.pdf"'
    return response

# -------- CIRCULAR DOCX --------
def download_circular_docx(request):
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    data = request.session.get("circular_data")
    if not data:
        return HttpResponse("No circular generated", status=400)

    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add BISAG Logo
    logo_path = os.path.join(settings.BASE_DIR, "static", "generator", "bisag_logo.png")
    if os.path.exists(logo_path):
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, height=Inches(0.9))
        doc.add_paragraph()  # Add space after logo

    # Header lines
    for line in data["header"].values():
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

    # Circular title
    lang = data.get("language", "en")
    title_text = "परिपत्र" if lang == "hi" else "Circular"
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(16)
    
    # Date
    date_label = "दिनांक :" if lang == "hi" else "Date :"
    p = doc.add_paragraph(f"{date_label} {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Subject
    subject_label = "विषय :" if lang == "hi" else "Subject :"
    p = doc.add_paragraph(f"{subject_label} {data['subject']}")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Body
    p = doc.add_paragraph(data["body"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(12)

    # From section
    p = doc.add_paragraph(data["from"])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # Add some space before table
    doc.add_paragraph()
    
    # To section - Table
    to_people = data.get("to_people", [])
    if to_people:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        sr_label = "क्र." if lang == "hi" else "Sr. No."
        name_label = "नाम" if lang == "hi" else "Name"
        sign_label = "हस्ताक्षर" if lang == "hi" else "Sign"
        
        hdr_cells[0].text = sr_label
        hdr_cells[1].text = name_label
        hdr_cells[2].text = sign_label
        
        # Make header bold
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for idx, person in enumerate(to_people, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            name = person.get("name_hi") if lang == "hi" else person.get("name_en")
            row_cells[1].text = name or ""
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_cells[2].text = ""  # Empty for signature
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(3.5)
        table.columns[2].width = Inches(1.5)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="Circular.docx"'
    return response

# ===============================
# OFFICE ORDER FORM (MISSING FIX)
# ===============================
def office_order_form(request):
    return render(request, "generator/office_order_form.html", {
        "designations": DESIGNATION_MAP.keys()
    })

# =====================================================================
# ================= OFFICE ORDER PDF & DOCX (RESTORED) =================
# =====================================================================

def download_pdf(request):
    data = request.session.get("doc_data")
    if not data:
        return HttpResponse("No office order generated", status=400)

    html = render_to_string("generator/pdf_office_order.html", data)

    pdf = HTML(
        string=html,
        base_url=settings.BASE_DIR
    ).write_pdf(
        stylesheets=[
            CSS(string="""
                @page { size: A4; margin: 2.5cm; }
                body { font-family: serif; font-size: 12pt; line-height: 1.6; }
                .center { text-align: center; }
                .bold { font-weight: bold; }
                .ref-date-row { display: table; width: 100%; margin: 20px 0; }
                .ref-left { display: table-cell; text-align: left; font-weight: bold; width: 50%; }
                .date-right { display: table-cell; text-align: right; font-weight: bold; width: 50%; }
                .title { text-align: center; font-weight: bold; text-decoration: underline; margin: 20px 0; }
                .body { text-align: justify; margin: 20px 0; }
                .from-section { text-align: right; font-weight: bold; margin: 40px 0 20px; }
                .to-section { margin-top: 20px; }
                .to-section div { margin: 5px 0; }
            """)
        ]
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Office_Order.pdf"'
    return response


def download_docx(request):
    from docx.shared import Pt, Inches
    
    data = request.session.get("doc_data")
    if not data:
        return HttpResponse("No office order generated", status=400)

    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header
    for line in data["header"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

    # Reference & Date
    lang = data.get("language", "en")
    ref_label = "सं :" if lang == "hi" else "Ref:"
    p = doc.add_paragraph(f"{ref_label} {data['reference']}")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    date_label = "दिनांक :" if lang == "hi" else "Date:"
    p = doc.add_paragraph(f"{date_label} {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Title
    p = doc.add_paragraph(data["title"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(16)

    # Body
    p = doc.add_paragraph(data["body"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(12)

    # From
    p = doc.add_paragraph(data["from"])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # To
    to_label = "प्रति :" if lang == "hi" else "To:"
    p = doc.add_paragraph(to_label)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    for t in data["to"]:
        p = doc.add_paragraph(t)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="Office_Order.docx"'
        }
    )


# =====================================================================
# ===================== AUTHENTICATION VIEWS ==========================
# =====================================================================

def login_view(request):
    """User login view"""
    if request.user.is_authenticated:
        return redirect('home')
    
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            auth_login(request, user)
            return redirect('home')
        else:
            return render(request, 'generator/login.html', {
                'error': 'Invalid username or password. Please try again.'
            })
    
    return render(request, 'generator/login.html')


def register_view(request):
    """User registration view"""
    if request.user.is_authenticated:
        return redirect('home')
    
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        password1 = request.POST.get('password1')
        password2 = request.POST.get('password2')
        
        # Validation
        if password1 != password2:
            return render(request, 'generator/register.html', {
                'error': 'Passwords do not match. Please try again.'
            })
        
        if User.objects.filter(username=username).exists():
            return render(request, 'generator/register.html', {
                'error': 'Username already exists. Please choose a different one.'
            })
        
        if User.objects.filter(email=email).exists():
            return render(request, 'generator/register.html', {
                'error': 'Email already registered. Please use a different email.'
            })
        
        # Create user
        user = User.objects.create_user(
            username=username,
            email=email,
            password=password1,
            first_name=first_name,
            last_name=last_name
        )
        
        # Redirect to login page with success message
        return render(request, 'generator/login.html', {
            'success': 'Account created successfully! Please login with your credentials.'
        })
    
    return render(request, 'generator/register.html')


def logout_view(request):
    """User logout view"""
    auth_logout(request)
    return redirect('login')


# =====================================================================
# ======================== POLICY DOCUMENT ============================
# =====================================================================

def policy_form(request):
    """Policy form view"""
    return render(request, 'generator/policy_form.html', {
        "designations": DESIGNATION_MAP.keys()
    })


def generate_policy_title(request):
    """Generate policy subject with AI"""
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    prompt = request.POST.get("subject_prompt", "").strip()
    lang = request.POST.get("language", "en")

    if lang == "hi":
        system_prompt = """
आप BISAG-N के लिए एक सरकारी नीति दस्तावेज़ का विषय (Subject) लिख रहे हैं।

नियम:
- केवल एक पंक्ति का संक्षिप्त विषय लिखें।
- औपचारिक सरकारी भाषा का प्रयोग करें।
- 5-12 शब्दों में।
- केवल विषय लिखें, कुछ और नहीं।
"""
    else:
        system_prompt = """
You are writing the subject line for an official Government Policy document for BISAG-N.

Rules:
- Write ONE concise subject line only.
- Use formal government language.
- 5-12 words maximum.
- Return ONLY the subject, nothing else.
"""

    res = gemini_model.generate_content(system_prompt + "\n\nTopic:\n" + prompt)
    return HttpResponse(res.text.strip())


def generate_policy_body(request):
    """Generate policy body with AI"""
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    try:
        prompt = request.POST.get("body_prompt", "").strip()
        lang = request.POST.get("language", "en")
        
        if not prompt:
            return HttpResponse("Please provide a prompt for AI generation.", status=200)

        if lang == "hi":
            system_prompt = """
आप BISAG-N के लिए एक सरकारी नीति दस्तावेज़ की सामग्री लिख रहे हैं।

नियम:
- औपचारिक नीति भाषा का प्रयोग करें।
- उद्देश्य, दायरा, और प्रमुख बिंदुओं को शामिल करें।
- 3-5 अनुच्छेदों में लिखें।
- कोई शीर्षक या नीति संख्या न लिखें।
- सरकारी भाषा का प्रयोग करें।
- केवल सादा पाठ में उत्तर दें।
"""
        else:
            system_prompt = """
You are writing the content for an official Government Policy document for BISAG-N.

Rules:
- Use formal policy language.
- Include purpose, scope, and key points.
- Write 3-5 formal paragraphs.
- Do NOT include title or policy number.
- Official government tone.
- Plain text only.
"""

        res = gemini_model.generate_content(system_prompt + "\n\nTopic:\n" + prompt)
        generated_text = res.text.strip() if res and res.text else "Unable to generate content. Please try again."
        return HttpResponse(generated_text, status=200)
        
    except Exception as e:
        error_msg = f"AI generation failed: {str(e)}"
        print(f"[ERROR] Policy Gemini API: {error_msg}")
        return HttpResponse("Failed to generate content. Please check your internet connection and try again.", status=200)


def result_policy(request):
    """Policy preview"""
    if request.method != "POST":
        return redirect("policy_form")

    lang = request.POST.get("language")
    raw_date = request.POST.get("date")
    date = format_date_ddmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%m-%Y")
    
    # Get header based on language
    header = POLICY["header"][lang]
    
    # Handle multiple recipients
    to_recipients = request.POST.getlist("to_recipients[]")
    to_list = [DESIGNATION_MAP[pos][lang] for pos in to_recipients if pos in DESIGNATION_MAP]
    
    # Format recipients based on language (for PDF/DOCX download)
    if lang == "hi":
        to_str = ", ".join(to_list) if to_list else "सभी संबंधित"
    else:
        to_str = ", ".join(to_list) if to_list else "All Concerned"
    
    data = {
        "language": lang,
        "header": header,
        "date": date,
        "subject": request.POST.get("subject"),
        "body": request.POST.get("body"),
        "from": DESIGNATION_MAP[request.POST.get("from_position")][lang],
        "to": to_str,
        "to_list": to_list,  # Pass list for template to display one per line
    }

    request.session["policy_data"] = data
    return render(request, "generator/result_policy.html", data)


def download_policy_pdf(request):
    """Download policy as PDF"""
    data = request.session.get("policy_data")
    if not data:
        return HttpResponse("No policy generated", status=400)

    html = render_to_string("generator/pdf_policy.html", data)

    pdf = HTML(
        string=html,
        base_url=settings.BASE_DIR
    ).write_pdf(
        optimize_images=True,
        jpeg_quality=85,
        presentational_hints=True
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Policy.pdf"'
    return response


def download_policy_docx(request):
    """Download policy as DOCX"""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    data = request.session.get("policy_data")
    if not data:
        return HttpResponse("No policy generated", status=400)

    doc = Document()
    
    # Header
    for line in data["header"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # spacing
    
    # Title
    title = "नीति" if data["language"] == "hi" else "POLICY"
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)
    p.runs[0].underline = True
    
    doc.add_paragraph()  # spacing
    
    # Date
    date_label = "दिनांक:" if data["language"] == "hi" else "Date:"
    p = doc.add_paragraph(f"{date_label} {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    
    # Subject
    subject_label = "विषय:" if data["language"] == "hi" else "Subject:"
    p = doc.add_paragraph(f"{subject_label} {data['subject']}")
    p.runs[0].bold = True
    
    doc.add_paragraph()  # spacing
    
    # Body
    p = doc.add_paragraph(data["body"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # spacing
    
    # From (no label, right aligned)
    p = doc.add_paragraph(data['from'])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    
    # To (with recipients one per line)
    to_label = "प्रति:" if data["language"] == "hi" else "To:"
    p = doc.add_paragraph(to_label)
    p.runs[0].bold = True
    
    # Add recipients one per line
    if 'to_list' in data and data['to_list']:
        for recipient in data['to_list']:
            p = doc.add_paragraph(f"    {recipient}")  # indented
    else:
        p = doc.add_paragraph(f"    {data['to']}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="Policy.docx"'
        }
    )
